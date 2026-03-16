#!/usr/bin/env python3
"""Generate a professionally formatted DOCX changelog for Ah Ho Fruit handover."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(11.5)


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)


def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Ah Ho Fruit')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Changelog & Handover Document')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

meta_items = [
    ('Website:', 'ahhofruit.com'),
    ('Platform:', 'WordPress + WooCommerce (Avada Theme)'),
    ('Hosting:', 'Vodien (Singapore)'),
    ('Project Duration:', '21 January 2026 — 10 March 2026'),
    ('Total Commits:', '320'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document generated: 10 March 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# PROJECT OVERVIEW
# ============================================================
doc.add_heading('Project Overview', level=1)

add_para(
    'A full-scope e-commerce customisation project for Ah Ho Fruit, a Singapore-based wholesale '
    'and retail fruit supplier. The project involved building a complete B2B wholesale ordering '
    'system on top of WordPress/WooCommerce, with custom invoicing, delivery workflows, salesperson '
    'management, and inventory tools — all tailored for the fresh produce industry.'
)

doc.add_heading('Key Deliverables', level=2)
deliverables = [
    'Custom B2B wholesale pricing and ordering system',
    'PDF Invoice, Delivery Order, and Packing Slip generation',
    'Salesperson and Storeman role management',
    'WhatsApp B2B catalog generator',
    'Quick Stock Update tool for bulk inventory management',
    'Delivery date scheduling with business rules',
    'Payment gateway fee processing',
    'GST handling for B2B orders',
    'Mobile app access for field staff',
    'Comprehensive documentation and guides',
]
for d in deliverables:
    add_bullet(d)

add_hr()

# ============================================================
# WEEK 1-2
# ============================================================
doc.add_heading('Week 1–2: 21 Jan – 31 Jan 2026', level=1)

doc.add_heading('Project Setup & Infrastructure (21 Jan)', level=2)
add_bullet(' — Set up automated deployment from GitHub to Vodien hosting using FTP (GitHub Actions). Tested multiple deployment methods (rsync, SFTP) before settling on FTP due to Vodien\'s hosting restrictions.', 'WordPress deployment pipeline created')
add_bullet(' — Licensed premium WordPress theme (ThemeForest) installed and configured as the storefront foundation.', 'Avada theme installed')
add_bullet(' — Full project repository set up on GitHub for code tracking and automated deployments.', 'Git version control initialised')

doc.add_heading('Custom Order Statuses (23 Jan)', level=2)
add_bullet(' — Created custom WooCommerce order statuses to match Ah Ho Fruit\'s delivery workflow: "Out for Delivery", "Delivered", "Ready for Pickup".', 'Delivery workflow statuses added')
add_bullet(' — Automatic customer email notifications for each status change, so customers are informed when their order is out for delivery or ready for pickup.', 'Email notifications configured')

doc.add_heading('B2B Salesperson System (24 Jan)', level=2)
add_bullet(' — New WordPress user role specifically for Ah Ho Fruit\'s sales staff, with restricted access (can only see their own orders and customers).', 'Salesperson role created')
add_bullet(' — Automatic commission calculation for salesperson orders, with admin tools to view and manage commissions.', 'Commission tracking system built')
add_bullet(' — Salespersons can create orders and manage their customers but cannot access admin settings, other staff\'s orders, or financial reports.', 'Role-based access control')

doc.add_heading('PDF Invoicing System (24 Jan)', level=2)
add_para('Built a complete 4-phase PDF document system:')
add_bullet(' — Professional PDF invoices generated automatically for every order, featuring Ah Ho Fruit branding, itemised line items, weights, and totals.', 'Phase 1 — Invoice Generation')
add_bullet(' — Separate PDF documents for warehouse packing and delivery drivers. Packing slips show item details and weights; delivery orders match traditional commercial format with sign-off fields.', 'Phase 2 — Packing Slip & Delivery Order')
add_bullet(' — Invoices automatically attached to WooCommerce order confirmation emails. Configurable settings for which PDF types to auto-send.', 'Phase 3 — Email Automation')
add_bullet(' — Admin page to select multiple orders and download all their PDFs as a single ZIP file. Useful for monthly accounting or batch printing.', 'Phase 4 — Bulk PDF Download')

doc.add_heading('SEO Foundation (25 Jan)', level=2)
add_bullet(' — Meta tags, Open Graph tags, structured data (Schema.org), XML sitemap generation, and robots.txt configured for search engine visibility.', 'Comprehensive SEO setup')
add_bullet(' — Standardised SEO title and description templates created for all product categories.', 'Product SEO templates')

doc.add_heading('Server Stability & Troubleshooting (25 Jan)', level=2)
add_bullet(' — Resolved a server 500 error caused by .htaccess conflicts between PHP version requirements and Vodien\'s server configuration. Required extensive diagnosis across PHP versions, OPcache, and WordPress core files.', 'Critical site recovery')
add_bullet(' — Established stable FTP deployment path after thorough testing with Vodien\'s server structure.', 'Deployment pipeline refined')

doc.add_heading('Product Add-ons Plugin (25 Jan)', level=2)
add_bullet(' — Customers can add personal gift messages to their orders during checkout. Messages appear on packing slips with highlighted styling for warehouse staff.', 'Gift message support')
add_bullet(' — Additional notes field for special instructions (e.g., "extra ripe", "for event on Saturday").', 'Customer notes field')
add_bullet(' — Add-on indicators styled with dashed borders and clear markers for printing on standard black-and-white printers.', 'B&W printer optimisation')

doc.add_heading('Legal Compliance (26 Jan)', level=2)
add_bullet(' — Auto-generated legal page with e-commerce terms specific to fresh produce delivery (perishable goods policy, delivery timeframes, refund conditions).', 'Terms & Conditions page')
add_bullet(' — PDPA-compliant privacy policy covering data collection, storage, and customer rights.', 'Privacy Policy page')
add_bullet(' — Legal pages automatically linked in the website footer menu.', 'Footer integration')

doc.add_heading('Typography & Readability (26 Jan)', level=2)
add_bullet(' — Fixed font sizing, line height, and colour contrast across the Avada theme for better readability on all pages.', 'Site-wide readability improvements')

doc.add_heading('PHP 8 Compatibility Fixes (29 Jan)', level=2)
add_bullet(' — Resolved PHP 8.x type errors in WooCommerce checkout where product weight calculations were failing. Applied type-casting fixes across invoice, packing slip, and checkout processes.', 'Weight calculation fix')

add_hr()

# ============================================================
# WEEK 3-4
# ============================================================
doc.add_heading('Week 3–4: 1 Feb – 14 Feb 2026', level=1)

doc.add_heading('Payment Gateway Fees (31 Jan – 1 Feb)', level=2)
add_bullet(' — Built a plugin to automatically apply payment processing fees (e.g., 3.4% + $0.50 for Stripe credit card payments). Fees are transparently shown to customers at checkout.', 'Credit card surcharge system')
add_bullet(' — Configured PayNow (Singapore bank transfer) as the default payment method, encouraging lower-cost payment options.', 'PayNow default payment')
add_bullet(' — Ensured payment fee calculations work with both the classic checkout and the newer WooCommerce Blocks checkout.', 'WooCommerce Blocks compatibility')
add_bullet(' — Processing fees correctly calculated on the full order amount including shipping costs.', 'Fee calculation includes shipping')

doc.add_heading('B2B Wholesale Pricing System (1 Feb)', level=2)
add_bullet(' — Added a separate wholesale price field to every WooCommerce product. When a salesperson creates an order, the wholesale price is automatically applied instead of the retail price.', 'Wholesale price field')
add_bullet(' — System automatically identifies B2B orders based on whether products have wholesale pricing applied, ensuring correct pricing on invoices and reports.', 'Price detection logic')
add_bullet(' — All wholesale pricing features built to work with WooCommerce\'s High-Performance Order Storage (modern database structure).', 'HPOS compatibility')

doc.add_heading('Salesperson Enhancements (1–2 Feb)', level=2)
add_bullet(' — Salespersons can only view and edit their own orders. Order counts and status badges in the admin menu reflect only their assigned orders.', 'Order restriction')
add_bullet(' — Salespersons can create new customer accounts and set payment terms (e.g., "COD", "7 Days", "14 Days", "30 Days").', 'Customer management')
add_bullet(' — Removed unnecessary WordPress admin menu items for salesperson accounts, showing only relevant sections (Orders, Customers, Products).', 'Admin menu cleanup')
add_bullet(' — Salespersons are automatically redirected to the Orders page after login, not the WordPress dashboard.', 'Login redirect')
add_bullet(' — Enabled product search capability when salespersons create orders.', 'Product search in orders')
add_bullet(' — Every order tracks which salesperson created it, displayed as a column in the orders list and in order details.', 'Salesperson attribution')
add_bullet(' — Added a one-click "Complete" button for B2B orders in the orders list, streamlining the fulfilment workflow.', 'Quick "Complete" action')

doc.add_heading('Commission System Upgrade (2 Feb)', level=2)
add_bullet(' — Commissions are now calculated the moment an order is placed, not delayed.', 'Immediate calculation')
add_bullet(' — Admin tool to recalculate commissions for existing orders (useful for retroactive adjustments).', 'Recalculation tool')

doc.add_heading('Custom Order Status — Processing B2B (2 Feb)', level=2)
add_bullet(' — New order status specifically for wholesale/B2B orders. Automatically assigned when a salesperson creates an order.', '"Processing - B2B" status')

doc.add_heading('Customer Payment Terms (2 Feb)', level=2)
add_bullet(' — Added configurable payment terms to customer profiles (COD, 7 Days, 14 Days, 30 Days, or custom terms).', 'Payment terms field')
add_bullet(' — Payment terms can be changed directly from the order page without reloading.', 'Dynamic AJAX update')
add_bullet(' — Configurable list of payment term options in the plugin settings.', 'Admin settings UI')

doc.add_heading('WhatsApp B2B Catalog Generator (2 Feb)', level=2)
add_bullet(' — Admin page that generates a WhatsApp-friendly text catalog of all B2B products with wholesale prices. Staff can copy and paste directly into WhatsApp.', 'One-click catalog generation')
add_bullet(' — Only shows in-stock products that have a wholesale price set. Excludes B2C-only products.', 'Smart filtering')
add_bullet(' — Products organised by category with clear headings.', 'Category grouping')
add_bullet(' — Catalog updates automatically whenever prices or stock levels change.', 'Auto-refresh')

doc.add_heading('B2B Stock List (5 Feb)', level=2)
add_bullet(' — Added a separate "B2B Stock List" section showing stock quantities alongside prices. For internal reference only — copy protection enabled.', 'Internal stock reference')

doc.add_heading('Delivery Date System (2 Feb)', level=2)
add_bullet(' — Customers select their preferred delivery date at checkout using a visual calendar (Flatpickr). Sundays greyed out. Public holidays blocked.', 'Checkout date picker')
add_bullet(' — Minimum lead time enforced (next business day). Saturday deliveries allowed. Sunday deliveries blocked.', 'Business rules enforced')
add_bullet(' — Staff can view and modify the delivery date from the order admin page.', 'Admin editing')

doc.add_heading('Delivery Order Redesign (6–7 Feb)', level=2)
add_bullet(' — Redesigned the Delivery Order PDF to match the traditional Singapore commercial format, with proper header layout, company details (UEN, GST registration), and sign-off section.', 'Traditional commercial format')
add_bullet(' — Tightened margins and layout to ensure most delivery orders fit on a single A4 page.', 'Single-page optimisation')
add_bullet(' — PO Number and Payment Terms automatically populated from order data.', 'Auto-fill fields')
add_bullet(' — WhatsApp number and business email displayed on the delivery order.', 'Contact details')

doc.add_heading('Dual Commission Model (7 Feb)', level=2)
add_bullet(' — Extended the commission system to support both flat per-carton rates and percentage-based commissions.', 'Per-carton + percentage commissions')

doc.add_heading('PDF Improvements (7 Feb)', level=2)
add_bullet(' — Added print buttons to invoice, delivery order, and packing slip pages.', 'Print buttons')
add_bullet(' — Optimised PDF layouts with tighter margins. Standardised delivery orders with 14 line item rows.', 'Layout optimisation')
add_bullet(' — Gift messages and customer requests highlighted with coloured boxes on packing slips.', 'Gift/request highlighting')

doc.add_heading('Storeman Role (8 Feb)', level=2)
add_bullet(' — Created a dedicated role for warehouse staff with access to inventory management but restricted from financial and order management features.', 'New "Storeman" role')
add_bullet(' — Storemen can view and update stock levels but cannot modify prices or access customer data.', 'Product inventory access')
add_bullet(' — Staff members can only see customers they personally created.', 'Customer visibility restriction')

doc.add_heading('Invoice Redesign (8 Feb)', level=2)
add_bullet(' — Redesigned the invoice PDF template to match the delivery order\'s professional layout, ensuring consistent branding.', 'Matching layout')
add_bullet(' — Simplified invoice numbering to use the WooCommerce order number directly.', 'Order number as invoice number')

doc.add_heading('Quick Stock Update Page (8 Feb)', level=2)
add_bullet(' — Full-featured admin page for bulk stock management, allowing warehouse staff to update stock quantities for all products in one screen.', 'Dedicated inventory page')
add_bullet(' — Filter products by category for faster stock-taking.', 'Category filtering')
add_bullet(' — Quick search to find specific products.', 'Search functionality')
add_bullet(' — Visual indicators (yellow highlight) showing which rows have been modified before saving.', 'Change tracking')
add_bullet(' — Optimised layout for tablet use during warehouse stock-taking.', 'Mobile responsive')

doc.add_heading('Additional Features (8–9 Feb)', level=2)
add_bullet(' — Installed the HeyMag Chat widget for customer support automation on the website.', 'HeyMag Chat Plugin')
add_bullet(' — Made wholesale products visible in the shop for retail customers too, expanding product discoverability.', 'B2B Product Visibility')
add_bullet(' — Added Noto Sans CJK font for Chinese character support in all PDFs.', 'CJK Font Support')
add_bullet(' — Redesigned packing slips for paper efficiency. Removed sign-off section (kept on delivery orders only).', 'Packing Slip Optimisation')

doc.add_heading('Omakase Boxes Promotion (12 Feb)', level=2)
add_bullet(' — Pinned "Omakase Boxes" products to the top of the shop page for promotional visibility.', 'Product pinning')

doc.add_heading('Salesman Account Setup (13 Feb)', level=2)
add_bullet(' — Set up 5 salesman accounts with proper roles, capabilities, and access restrictions.', 'Bulk account creation')

doc.add_heading('Partial Delivery & Returns System (14 Feb)', level=2)
add_bullet(' — Built a system to track partial deliveries for large B2B orders. Each delivery generates its own PDF delivery order.', 'Partial delivery tracking')
add_bullet(' — Returns tracking within orders, showing returned items and adjusted totals on invoices.', 'Item returns handling')
add_bullet(' — All delivery and return actions logged with timestamps and user attribution.', 'Audit trail')

add_hr()

# ============================================================
# WEEK 5-6
# ============================================================
doc.add_heading('Week 5–6: 15 Feb – 28 Feb 2026', level=1)

doc.add_heading('FTP Deployment Fix (15 Feb)', level=2)
add_bullet(' — Discovered and fixed a critical deployment path issue where the server directory was "ah-ho-fruit" (no "s"), not "ah-ho-fruits". Extensive diagnosis required.', 'Server path correction')
add_bullet(' — Cleaned up the incorrectly-named directory on the server.', 'Stale directory cleanup')

doc.add_heading('Salesperson Meta Box Upgrade (15–16 Feb)', level=2)
add_bullet(' — Rebuilt the salesperson assignment meta box on the order page for HPOS compatibility.', 'HPOS-compatible assignment')
add_bullet(' — Salesperson assignment now saves via AJAX without requiring a full page reload.', 'AJAX save')
add_bullet(' — Commissions are correctly calculated even when an admin manually assigns a salesperson.', 'Commission on admin assignment')

doc.add_heading('Business Rebranding (19 Feb)', level=2)
add_bullet(' — Comprehensive rebrand across the entire codebase, changing all instances of "Ah Ho Fruits" (with "s") to "Ah Ho Fruit" (without "s"). Applied to plugin names, PDF templates, email templates, admin interface labels, WhatsApp catalog headers, and documentation.', '"Ah Ho Fruits" → "Ah Ho Fruit"')

doc.add_heading('Shop & Checkout Improvements (23 Feb)', level=2)
add_bullet(' — Customised shop page product sorting for logical, category-aware ordering.', 'Product ordering')
add_bullet(' — Added free delivery for orders above $60. Self-pickup remains available regardless of order amount.', 'Free shipping threshold')
add_bullet(' — Updated business contact email across all customer-facing templates.', 'Email update')

doc.add_heading('PDF Quality & Download Fixes (23 Feb)', level=2)
add_bullet(' — Resolved persistent Chinese character display issues by switching to pre-built font metrics and Medium weight Noto Sans CJK font. Characters that previously showed as "?????" now render correctly.', 'CJK font fix (final)')
add_bullet(' — Increased text darkness across all PDFs for better readability when printed on standard office printers.', 'Darker text for printing')
add_bullet(' — Implemented automatic PDF cache directory creation with graceful fallback on write failure.', 'PDF cache system')
add_bullet(' — Fixed Vodien\'s server proxy replacing PDF filenames with random UUIDs. Implemented workaround to preserve filenames like "Invoice-1234.pdf".', 'Download filename fix')

doc.add_heading('WooCommerce Mobile App Access (23 Feb)', level=2)
add_bullet(' — Enabled the WooCommerce mobile app for Storeman accounts, allowing warehouse staff to check and update stock levels from their phones.', 'Storeman app access')
add_bullet(' — Enabled the WooCommerce mobile app for Salesperson accounts, allowing field sales staff to create orders on the go.', 'Salesperson app access')

doc.add_heading('Order Item Editing & Returns (26 Feb)', level=2)
add_bullet(' — Built a system allowing admin staff to edit order line items (change quantities, add/remove products) after the order has been placed. All changes are logged.', 'Edit items after order creation')
add_bullet(' — Invoices now show return details and net totals when items have been returned.', 'Returns on invoice')

doc.add_heading('Documentation (25 Feb – 1 Mar)', level=2)
add_bullet(' — Updated comprehensive system documentation covering all custom plugins and features.', 'System documentation refresh')
add_bullet(' — Created a comprehensive Order Management Guide with step-by-step instructions for non-technical staff.', 'Order Management Guide')

add_hr()

# ============================================================
# WEEK 7-8
# ============================================================
doc.add_heading('Week 7–8: 1 Mar – 10 Mar 2026', level=1)

doc.add_heading('Product Add-on: Omakase Flowers (2 Mar)', level=2)
add_bullet(' — Added an optional "Add Flowers" checkbox for Omakase Box products. When selected, a bouquet is added at an additional charge.', 'Optional flower add-on')
add_bullet(' — Clicking the flower add-on shows a lightbox preview of the flower arrangement.', 'Lightbox preview')

doc.add_heading('Promotional Content Update (3 Mar)', level=2)
add_bullet(' — Changed promotional discount from "10% off" to "5% off" on product pages.', 'Discount update')
add_bullet(' — Changed "next morning delivery" to "next day delivery" to accurately reflect delivery timeframes.', 'Delivery promise update')

doc.add_heading('Invoice Deliver To Section (4 Mar)', level=2)
add_bullet(' — Added a "Deliver To" section alongside the "Bill To" section on invoices, showing the delivery address and phone number separately. Essential for B2B orders where billing and delivery addresses often differ.', 'Dual address layout')

doc.add_heading('Product Image Generation (4 Mar)', level=2)
add_bullet(' — Generated and uploaded a professional product image for the White-flesh Pomelo product using AI image generation. Iterated through multiple versions to achieve the correct appearance.', 'Pomelo product image')

doc.add_heading('Delivery Date Improvements (7 Mar)', level=2)
add_bullet(' — Fixed delivery date not saving correctly on the WooCommerce Blocks checkout page.', 'Blocks checkout save fix')
add_bullet(' — Added the selected delivery date to the invoice PDF template.', 'Invoice delivery date')
add_bullet(' — Fixed the delivery date edit fields not appearing when clicking "Edit" on the order admin page.', 'Admin edit fix')
add_bullet(' — Updated business rules to allow Saturday deliveries. Added Singapore public holiday blocking.', 'Saturday delivery & public holidays')
add_bullet(' — Simplified checkout by removing time slot selection that was not operationally supported.', 'Time slot removal')

doc.add_heading('Payment Fee Fix (7 Mar)', level=2)
add_bullet(' — Fixed credit card processing fees not being applied on Blocks checkout for guest customer orders.', 'Guest checkout fee fix')

doc.add_heading('Knowledge Base (7 Mar)', level=2)
add_bullet(' — Created a comprehensive knowledge base document covering all Ah Ho Fruit products, delivery policies, ordering processes, and FAQs. Powers the AI chat widget\'s responses.', 'HeyMag Copilot knowledge base')

doc.add_heading('B2B GST Implementation (7–8 Mar)', level=2)
add_bullet(' — Added a checkbox on the WooCommerce order page: "B2B Order — Add 9% GST to invoice". When checked, GST is added to the order.', 'GST toggle on orders')
add_bullet(' — The GST checkbox is automatically checked when a salesperson or storeman creates an order with wholesale-priced items.', 'Auto-enable for B2B')
add_bullet(' — Any admin or staff member can manually toggle the GST checkbox for flexibility.', 'Manual override')
add_bullet(' — GST (9%) amount and Total + GST rows displayed directly on the order admin page with real-time updates.', 'GST on admin order page')
add_bullet(' — B2B invoices show the GST amount and GST-inclusive total. B2C invoices remain unchanged (nett prices).', 'GST on invoice PDF')

doc.add_heading('WhatsApp Catalog Emoji Headings (8 Mar)', level=2)
add_para('Updated the WhatsApp B2B catalog to use dual emojis for each category heading:')
emoji_items = [
    '🍏 APPLES 🍎', '🍓 BERRIES 🫐', '🍊 CITRUS 🍋', '🍇 GRAPES 🍇', '🥝 KIWI 🥝',
    '🍈 MELONS 🍉', '🥕 OTHERS 🌴', '🍐 PEARS 🍐', '🍑 STONE 🍑', '🥭 TROPICAL 🍌',
]
for item in emoji_items:
    add_bullet(item)

doc.add_heading('Quick Stock Update — Price Editing (8 Mar)', level=2)
add_bullet(' — Added editable Regular Price, Sale Price, and Wholesale Price columns to the Quick Stock Update page (admin-only).', 'Bulk price editing')
add_bullet(' — Administrators can now update stock levels AND prices for all products on a single page.', 'Combined stock & price management')
add_bullet(' — Price columns are only visible to administrator accounts. Storeman accounts see only stock quantity fields.', 'Role restriction')

doc.add_heading('Quick Stock Update — Role Compatibility Fix (9 Mar)', level=2)
add_bullet(' — Fixed an issue where certain administrator accounts (e.g., "ahhofruit") could not see the price editing columns. The system was checking for the manage_options capability, which can fail when an account\'s role was changed after creation. Now checks the user\'s actual role assignment directly, ensuring all Administrator accounts can view and edit prices regardless of how the account was originally set up.', 'Administrator role detection fix')

doc.add_heading('Credit Card Fee & Express Checkout Fix (10 Mar)', level=2)
add_bullet(' — Changed the default payment method from PayNow to Stripe (Credit Card). The 3.5% processing fee now displays immediately on the cart page and checkout page, so customers always see the fee upfront before paying.', 'Default payment switched to Credit Card')
add_bullet(' — Apple Pay, Google Pay, and Link (express checkout buttons) are powered by Stripe. With Stripe as the default, the processing fee is already included in the total when customers use express checkout. Previously, customers could complete payment via Apple Pay or Google Pay without being charged the 3.5% processing fee.', 'Express checkout fee fix')
add_bullet(' — The processing fee now appears on the cart page (not just checkout), giving customers full price transparency from the moment they view their cart.', 'Fee visible on cart page')
add_bullet(' — Customers who proactively select PayNow (bank transfer) as their payment method see the processing fee removed, paying $0 in fees. This incentivises lower-cost payment methods.', 'PayNow = no fee')
add_bullet(' — Apple Pay, Google Pay, and Link buttons are hidden on the cart page to prevent premature checkout. They remain available on the checkout page where the fee is visible.', 'Express checkout disabled on cart page')
add_bullet(' — Apple Pay, Google Pay, and Link buttons are also hidden on individual product pages, directing customers through the standard Add to Cart → Checkout flow where fees are transparent.', 'Express checkout disabled on product pages')
add_bullet(' — A friendly banner displayed at the top of the cart and checkout pages: "Credit card payments come with a 3.5% fee, so PayNow is the happiest (and fee-free) way to pay." Encourages customers to choose PayNow for a fee-free experience.', 'PayNow tip banner')
add_bullet(' — The banner uses JavaScript injection to work with WooCommerce Blocks (React-based) cart and checkout pages, which don\'t support traditional PHP hooks. Retries briefly to handle asynchronous Blocks rendering.', 'PayNow tip banner (Blocks compatible)')
add_bullet(' — Added intelligent duplicate fee detection to prevent double-charging. The system now checks all configured gateway labels and fee amounts before adding a fee, preventing issues when Stripe\'s internal gateway IDs differ between session and order.', 'Duplicate fee prevention')
add_bullet(' — Added a fallback hook that checks every Stripe order at creation time. If the processing fee was missed during cart calculation, it is automatically added to the order before payment is finalised.', 'Safety net for all Stripe orders')

doc.add_heading('Delivery Date Improvements (10 Mar)', level=2)
add_bullet(' — Changed the minimum delivery lead time from 3 working days to the next available working day (Monday\u2013Saturday). Sundays and Singapore public holidays remain blocked.', 'Next working day delivery')
add_bullet(' — Orders placed before 12pm can select next working day delivery. Orders placed after 12pm must select at least 2 working days ahead. Applies to both classic checkout (PHP) and Blocks checkout (JavaScript).', '12pm cutoff rule')

add_hr()

# ============================================================
# COMPLETE FEATURE LIST TABLE
# ============================================================
doc.add_heading('Complete Feature List', level=1)

features = [
    ('1', 'Automated FTP Deployment', 'GitHub → Vodien FTP auto-deployment on every code push'),
    ('2', 'Avada Theme', 'Premium WordPress theme licensed and configured'),
    ('3', 'Custom Order Statuses', '"Out for Delivery", "Delivered", "Ready for Pickup", "Processing - B2B"'),
    ('4', 'Order Status Emails', 'Automatic customer notifications on status changes'),
    ('5', 'Salesperson Role', 'Custom role with restricted access, order filtering, and commission tracking'),
    ('6', 'Storeman Role', 'Warehouse staff role with inventory-only access'),
    ('7', 'Commission System', 'Dual model — per-carton flat rate + percentage-based commissions'),
    ('8', 'B2B Wholesale Pricing', 'Separate wholesale price field with automatic application for B2B orders'),
    ('9', 'PDF Invoice', 'Professional branded invoice with itemised details, weights, GST'),
    ('10', 'PDF Delivery Order', 'Traditional commercial format with sign-off, PO, terms'),
    ('11', 'PDF Packing Slip', 'Warehouse picking document with gift/note highlighting'),
    ('12', 'Consolidated Packing Slip', 'Multi-order combined packing slip for batch fulfilment'),
    ('13', 'Bulk PDF Download', 'ZIP download of multiple order PDFs'),
    ('14', 'CJK Font Support', 'Chinese characters render correctly on all PDFs'),
    ('15', 'WhatsApp Catalog Generator', 'One-click B2B product catalog for WhatsApp sharing'),
    ('16', 'B2B Stock List', 'Internal stock reference with copy protection'),
    ('17', 'Quick Stock Update', 'Bulk inventory management page with category filters'),
    ('18', 'Bulk Price Editing', 'Admin-only Regular, Sale, and Wholesale price editing'),
    ('19', 'Delivery Date Picker', 'Checkout calendar with business rules, weekend/holiday blocking'),
    ('20', 'Payment Gateway Fees', '3.5% credit card fee shown on cart + checkout, with express checkout coverage'),
    ('21', 'Credit Card Default', 'Stripe (Credit Card) as default; PayNow available with $0 fees'),
    ('22', 'Free Shipping ($60+)', 'Free delivery threshold with self-pickup always available'),
    ('23', 'Product Add-ons', 'Gift messages, customer notes, and optional flower add-on'),
    ('24', 'Partial Delivery Tracking', 'Track partial deliveries with per-delivery PDFs'),
    ('25', 'Item Returns Handling', 'Returns tracking with adjusted invoice totals'),
    ('26', 'Order Item Editing', 'Post-creation line item editing with audit trail'),
    ('27', 'B2B GST System', 'Toggle-based 9% GST for wholesale orders'),
    ('28', 'Customer Payment Terms', 'Configurable payment terms (COD, 7/14/30 Days)'),
    ('29', 'Omakase Product Pinning', 'Featured products pinned to top of shop page'),
    ('30', 'SEO Foundation', 'Meta tags, structured data, XML sitemap'),
    ('31', 'Legal Pages', 'Terms & Conditions + Privacy Policy (auto-generated)'),
    ('32', 'Typography Fixes', 'Site-wide readability improvements'),
    ('33', 'HeyMag Chat Widget', 'AI-powered customer support chatbot'),
    ('34', 'Mobile App Access', 'WooCommerce mobile app for Storeman and Salesperson'),
    ('35', 'Customer Visibility Control', 'Staff see only their own customers'),
    ('36', 'Salesperson Auto-attribution', 'Orders automatically linked to creating salesperson'),
    ('37', 'Product Image Generation', 'AI-generated product images (Pomelo)'),
    ('38', 'Business Rebranding', '"Ah Ho Fruits" → "Ah Ho Fruit" across all systems'),
]
add_table(['#', 'Feature', 'Description'], features)

# ============================================================
# CUSTOM PLUGINS TABLE
# ============================================================
doc.add_heading('Custom Plugins Delivered', level=1)

plugins = [
    ('ah-ho-custom', 'Core business logic — order statuses, salesperson roles, wholesale pricing, catalog generator, inventory management, delivery dates, commissions, B2B GST', '20 PHP modules'),
    ('ah-ho-invoicing', 'PDF generation — invoices, delivery orders, packing slips, email automation, bulk downloads', '5 templates + Dompdf'),
    ('ah-ho-product-addons', 'Product add-ons — gift messages, customer notes, flower add-on with lightbox', '3 PHP files'),
    ('payment-gateway-fees', 'Payment processing fee calculation and display', '1 PHP file'),
    ('ah-ho-legal-pages-setup', 'Auto-generated Terms & Privacy Policy', '1 PHP file'),
    ('ah-ho-typography-fix', 'Site-wide readability corrections', '1 PHP file'),
    ('heymag-chat', 'AI customer support chatbot widget', '1 PHP file'),
]
add_table(['Plugin', 'Purpose', 'Key Files'], plugins)

# ============================================================
# TECHNICAL INFRASTRUCTURE TABLE
# ============================================================
doc.add_heading('Technical Infrastructure', level=1)

infra = [
    ('CMS', 'WordPress 6.x'),
    ('E-commerce', 'WooCommerce 9.x (HPOS enabled)'),
    ('Theme', 'Avada (ThemeForest license)'),
    ('Hosting', 'Vodien (Singapore, shared hosting)'),
    ('PHP Version', '8.3'),
    ('Database', 'MySQL (prefix: wpgr_)'),
    ('Deployment', 'GitHub Actions → FTP (SamKirkland/FTP-Deploy-Action)'),
    ('PDF Engine', 'Dompdf 2.x'),
    ('CJK Font', 'Noto Sans CJK Medium'),
    ('Checkout', 'WooCommerce Blocks + Classic fallback'),
    ('Payment', 'Stripe (Credit Card) + PayNow (Bank Transfer)'),
]
add_table(['Component', 'Details'], infra)

# ============================================================
# FOOTER
# ============================================================
add_hr()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Total project scope: 320 commits, 7 custom plugins, 38 features delivered')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

# Save
output_path = os.path.join(os.path.dirname(__file__), 'CHANGELOG-HANDOVER.docx')
doc.save(output_path)
print(f"Saved to {output_path}")
